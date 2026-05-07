"""Resume parser.

Two-stage pipeline:
1. Extract raw text from a file-shaped input (.pdf or .docx).
2. Run a structured-extraction LLM call to lift fields into a
   :class:`ParsedResume` (mirrors the candidate schema).

Stage 1 is deterministic and offline. Stage 2 uses the lightweight tier of
the model router (Haiku) with a JSON-mode prompt so it returns a strict
schema.
"""

from __future__ import annotations

import io
import json
import re
from datetime import date
from typing import Any, Protocol

import pypdf
import structlog
from docx import Document
from pydantic import BaseModel, ConfigDict, Field

from wfi_llm import ModelRouter, ModelTier
from wfi_schemas import (
    CareerHistoryEntry,
    Citizenship,
    ClearanceType,
    CompensationEntry,
    PolygraphType,
    SalesMotion,
    SeOrientation,
)

log = structlog.get_logger("wfi.data.resume")


class ParsedResume(BaseModel):
    """Structured resume extraction. Maps directly onto Candidate fields."""

    model_config = ConfigDict(use_enum_values=True)

    first_name: str = ""
    last_name: str = ""
    email: str | None = None
    phone: str | None = None
    linkedin_url: str | None = None
    location_city: str | None = None
    location_state: str | None = None
    location_metro: str | None = None
    citizenship: Citizenship = Citizenship.UNKNOWN
    clearance_type: ClearanceType = ClearanceType.NONE
    polygraph: PolygraphType = PolygraphType.NONE

    summary: str = ""
    skills: list[str] = Field(default_factory=list)
    education: list[dict[str, Any]] = Field(default_factory=list)
    career_history: list[CareerHistoryEntry] = Field(default_factory=list)
    compensation_history: list[CompensationEntry] = Field(default_factory=list)
    primary_motion: SalesMotion | None = None
    se_orientation: SeOrientation | None = None
    methodology_experience: list[str] = Field(default_factory=list)

    raw_text: str = ""
    extraction_confidence: float = Field(default=0.0, ge=0.0, le=1.0)

    def field_count(self) -> int:
        """Count of non-empty fields — used to score parser accuracy."""
        non_empty = 0
        for name, value in self.model_dump().items():
            if name == "raw_text":
                continue
            if isinstance(value, str | list | dict) and not value:
                continue
            if value is None:
                continue
            non_empty += 1
        return non_empty


# --- text extraction --------------------------------------------------------

def parse_pdf_bytes(data: bytes) -> str:
    """Extract text from a PDF byte string. Returns empty string on failure."""
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        log.warning("pdf_open_failed", error=str(exc))
        return ""
    chunks: list[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            log.warning("pdf_page_failed", error=str(exc))
    return "\n".join(chunks).strip()


def parse_docx_bytes(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        log.warning("docx_open_failed", error=str(exc))
        return ""
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def extract_text_from_bytes(data: bytes, *, filename: str = "") -> str:
    lower = filename.lower()
    if lower.endswith(".pdf") or data[:4] == b"%PDF":
        return parse_pdf_bytes(data)
    if lower.endswith(".docx") or data[:2] == b"PK":
        return parse_docx_bytes(data)
    # Fallback: assume utf-8 text.
    try:
        return data.decode("utf-8", errors="ignore").strip()
    except Exception:  # noqa: BLE001
        return ""


# --- regex pre-extraction ---------------------------------------------------

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(
    r"(?:\+?1[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}"
)
LINKEDIN_RE = re.compile(r"https?://(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+")

CLEARANCE_PATTERNS: list[tuple[re.Pattern, ClearanceType]] = [
    (re.compile(r"\bts/?sci\b|\btop secret(?:[/\\]| )sci\b", re.I), ClearanceType.TS_SCI),
    (re.compile(r"\btop secret\b|\bts\b", re.I), ClearanceType.TOP_SECRET),
    (re.compile(r"\bsecret\b", re.I), ClearanceType.SECRET),
    (re.compile(r"\bpublic trust\b", re.I), ClearanceType.PUBLIC_TRUST),
]

POLY_PATTERNS: list[tuple[re.Pattern, PolygraphType]] = [
    (re.compile(r"\blifestyle (?:poly|polygraph)\b", re.I), PolygraphType.LIFESTYLE),
    (re.compile(r"\bfull[- ]?scope (?:poly|polygraph)\b", re.I), PolygraphType.FULL_SCOPE),
    (re.compile(r"\bci (?:poly|polygraph)\b|\bcounter[- ]?intelligence poly", re.I), PolygraphType.CI),
]


def _regex_pre_extract(text: str) -> dict[str, Any]:
    out: dict[str, Any] = {}
    if (m := EMAIL_RE.search(text)):
        out["email"] = m.group(0).lower()
    if (m := PHONE_RE.search(text)):
        out["phone"] = re.sub(r"[^\d+]", "", m.group(0))
    if (m := LINKEDIN_RE.search(text)):
        out["linkedin_url"] = m.group(0)
    for pattern, value in CLEARANCE_PATTERNS:
        if pattern.search(text):
            out["clearance_type"] = value.value
            break
    for pattern, value in POLY_PATTERNS:
        if pattern.search(text):
            out["polygraph"] = value.value
            break
    return out


# --- LLM-driven structured extraction ---------------------------------------

class StructuredExtractor(Protocol):
    async def extract(self, text: str, *, hints: dict[str, Any]) -> dict[str, Any]: ...


_SYSTEM_PROMPT = """\
You are a resume parser for a recruiting platform. Extract structured data
from the provided resume text. Return STRICT JSON only — no commentary, no
markdown, no preamble.

Schema:
{
  "first_name": string,
  "last_name": string,
  "email": string|null,
  "phone": string|null,
  "linkedin_url": string|null,
  "location_city": string|null,
  "location_state": string|null,
  "location_metro": string|null,
  "citizenship": "us_citizen"|"permanent_resident"|"visa_h1b"|"visa_other"|"unknown",
  "clearance_type": "none"|"public_trust"|"secret"|"top_secret"|"ts_sci",
  "polygraph": "none"|"ci"|"full_scope"|"lifestyle",
  "summary": string,
  "skills": [string],
  "education": [{"school": string, "degree": string|null, "field": string|null, "year": int|null}],
  "career_history": [
    {
      "company": string,
      "title": string,
      "start_date": string|null,
      "end_date": string|null,
      "company_tier": "platform"|"established"|"growth"|"early"|null,
      "motion_type": "enterprise"|"mid_market"|"smb_velocity"|"plg"|"channel"|null,
      "quota_level": number|null,
      "territory": string|null,
      "notes": string|null
    }
  ],
  "compensation_history": [
    {"role_id": string|null, "base": number|null, "ote": number|null,
     "variable_structure": string|null, "year": int|null}
  ],
  "primary_motion": "enterprise"|"mid_market"|"smb_velocity"|"plg"|"channel"|null,
  "se_orientation": "pre_sales"|"post_sales"|"hybrid"|null,
  "methodology_experience": [string],
  "extraction_confidence": number  // 0.0 - 1.0
}

If a field is unknown, use null for scalars or [] for lists. Dates must be
ISO 8601 (YYYY-MM-DD); use the first of the month if only month/year is
known. Never invent compensation; only include amounts the resume explicitly
mentions and set w2_verified to false.
"""


class ResumeParser:
    def __init__(self, router: ModelRouter | None = None) -> None:
        self._router = router

    async def parse(self, raw_text: str) -> ParsedResume:
        if not raw_text.strip():
            return ParsedResume(raw_text="", extraction_confidence=0.0)

        regex_hints = _regex_pre_extract(raw_text)
        llm_data = await self._extract_with_llm(raw_text, hints=regex_hints)
        merged = {**regex_hints, **llm_data}
        merged["raw_text"] = raw_text

        # Coerce date strings to date objects for the typed fields.
        for entry in merged.get("career_history", []) or []:
            for key in ("start_date", "end_date"):
                value = entry.get(key)
                if isinstance(value, str):
                    entry[key] = _parse_date(value)
        # Coerce comp years.
        for entry in merged.get("compensation_history", []) or []:
            year = entry.get("year")
            if isinstance(year, str) and year.isdigit():
                entry["year"] = int(year)

        return ParsedResume(**merged)

    async def _extract_with_llm(
        self,
        text: str,
        *,
        hints: dict[str, Any],
    ) -> dict[str, Any]:
        if self._router is None:
            # No LLM configured — return whatever the regex pre-pass found.
            return {"extraction_confidence": 0.4 if hints else 0.1}
        # Truncate very long resumes to stay within input budget.
        truncated = text[:12_000]
        user = (
            f"Pre-extracted hints (verify against the text): {json.dumps(hints)}\n\n"
            f"Resume text:\n```\n{truncated}\n```"
        )
        response = await self._router.acomplete(
            tier=ModelTier.LIGHT,
            system=_SYSTEM_PROMPT,
            user=user,
            max_tokens=4096,
            temperature=0.0,
        )
        try:
            return json.loads(_strip_code_fence(response.text))
        except json.JSONDecodeError as exc:
            log.warning("resume_llm_json_parse_failed", error=str(exc), preview=response.text[:200])
            return {"extraction_confidence": 0.0}


def _strip_code_fence(value: str) -> str:
    stripped = value.strip()
    if stripped.startswith("```"):
        # remove first line (```json) and trailing ```
        lines = stripped.splitlines()
        if lines:
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        return "\n".join(lines).strip()
    return stripped


def _parse_date(value: str) -> date | None:
    if not value:
        return None
    formats = ("%Y-%m-%d", "%Y-%m", "%Y", "%m/%Y", "%m/%d/%Y")
    for fmt in formats:
        try:
            parsed = date.fromisoformat(value) if fmt == "%Y-%m-%d" else None
            if parsed is None:
                from datetime import datetime as dt
                parsed = dt.strptime(value, fmt).date()
            return parsed
        except ValueError:
            continue
    return None
